#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ScienceDirect 结果 → 最终报告
1. Crossref 补被引量（DOI 匹配）
2. DeepSeek 分类到 8 大类 + 深度校验（排除纯综述/纯概念）
3. 排序：优先今年 → 影响因子 × log10(被引+1)
4. 去重：累计记录已推送 DOI，每日不重复
5. 生成 Word/MD + Server酱微信提醒
"""
import json
import math
import os
import sys
import re
import datetime
import time

import requests

import common as fr

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

HERE = os.path.dirname(os.path.abspath(__file__))
SD_JSON = os.path.join(HERE, "sciencedirect_results.json")
STATE_PATH = os.path.join(HERE, "sd_state.json")
UA = "paper-report/1.0 (mailto:paper.report@example.com)"

CATEGORIES = {
    "1": "材料革新",
    "2": "界面粘结",
    "3": "约束混凝土",
    "4": "机器学习/深度学习",
    "5": "海上风电",
    "6": "多灾害韧性",
    "7": "LCA/数字孪生",
    "8": "新型组合体系",
}


def norm(s):
    return re.sub(r"[\s\-&]+", "", (s or "").lower())


def load_state():
    if os.path.exists(STATE_PATH):
        with open(STATE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"reported_dois": [], "last_run": None}


def save_state(state):
    with open(STATE_PATH, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


def crossref_citations(doi):
    try:
        r = requests.get(f"https://api.crossref.org/works/{doi}",
                         headers={"User-Agent": UA}, timeout=60)
        r.raise_for_status()
        return r.json().get("message", {}).get("is-referenced-by-count", 0) or 0
    except Exception:
        return 0


def extract_year(date_str):
    m = re.search(r"(\d{4})", date_str or "")
    return int(m.group(1)) if m else 0


def llm_classify(cfg, title, abstract):
    """返回 (category, depth_ok)。category 为 1-8 或 None。"""
    api = cfg["llm"]
    prompt = (
        "你是土木工程组合结构专家。判断论文属于以下哪个类别（输出数字 1-8，都不属于则输出 0），"
        "并判断是否包含实验验证或严谨数学推导（排除纯概念设想、以及缺少实验对比的综述）。\n"
        "1 材料革新：不锈钢/铝合金/双金属/FRP-钢复合管混凝土、UHPC、再生骨料/地聚物/橡胶/轻骨料等新型填充组合\n"
        "2 界面粘结：钢管/FRP筋-混凝土界面粘结滑移本构、传力机理、AE声发射/DIC监测、离散元/细观模拟\n"
        "3 约束混凝土：约束效率、FRP/螺旋箍/多腔约束、UHPC/ECC约束、SMA主动约束、损伤本构、非均匀约束\n"
        "4 机器学习/深度学习：承载力预测、本构曲面、LSTM/GAN/PINN/Transformer/CV、SHAP可解释性等\n"
        "5 海上风电：单桩/导管架/灌浆连接/吸力基础/疲劳/冲刷/腐蚀/冰荷载\n"
        "6 多灾害韧性：火灾-海水冷却、地震-海啸、冻融-疲劳耦合\n"
        "7 LCA/数字孪生：时变可靠度、碳排放约束、数字孪生、拓扑优化、材料-结构一体化设计\n"
        "8 新型组合体系：预应力FRP、可拆卸、3D打印、自复位、装配式、拱桥、框架-剪力墙\n"
        "只输出 JSON：{\"category\": 数字, \"depth\": true/false}"
    )
    resp = requests.post(
        api["base_url"].rstrip("/") + "/chat/completions",
        headers={"Authorization": f"Bearer {api['api_key']}", "Content-Type": "application/json"},
        json={
            "model": api["model"],
            "messages": [
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"标题：{title}\n摘要：{(abstract or '')[:2000]}"},
            ],
            "temperature": 0,
            "response_format": {"type": "json_object"},
        },
        timeout=90,
    )
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"].strip().strip("`")
    if content.lower().startswith("json"):
        content = content[4:].strip()
    data = json.loads(content)
    cat = str(data.get("category", 0))
    if cat not in CATEGORIES:
        cat = None
    return cat, bool(data.get("depth"))


def llm_analysis_with_intro(cfg, abstract, introduction):
    api = cfg["llm"]
    resp = requests.post(
        api["base_url"].rstrip("/") + "/chat/completions",
        headers={"Authorization": f"Bearer {api['api_key']}", "Content-Type": "application/json"},
        json={
            "model": api["model"],
            "messages": [
                {"role": "system",
                 "content": "你是土木工程学术助理。根据论文摘要与引言，输出中英双语严格JSON："
                            "{\"目的\":\"中文1-2句\",\"方法\":\"中文2-3句\",\"结论\":\"中文2-3句\","
                            "\"objective\":\"English 1-2 sentences\",\"methods\":\"English 2-3 sentences\","
                            "\"conclusions\":\"English 2-3 sentences\"}。只输出JSON。"},
                {"role": "user",
                 "content": f"摘要：{(abstract or '')[:2000]}\n\n引言：{(introduction or '')[:2000]}"},
            ],
            "temperature": 0.3,
            "response_format": {"type": "json_object"},
        },
        timeout=90,
    )
    resp.raise_for_status()
    content = resp.json()["choices"][0]["message"]["content"].strip().strip("`")
    if content.lower().startswith("json"):
        content = content[4:].strip()
    data = json.loads(content)
    return {
        "objective": data.get("目的", ""),
        "methods": data.get("方法", ""),
        "conclusions": data.get("结论", ""),
        "objective_en": data.get("objective", ""),
        "methods_en": data.get("methods", ""),
        "conclusions_en": data.get("conclusions", ""),
    }


def llm_summary(cfg, papers):
    """综合今日全部论文，生成「总结与思考」。"""
    items = []
    for i, p in enumerate(papers, 1):
        cat = CATEGORIES.get(p.get("category"), "其他")
        concl = p.get("analysis", {}).get("conclusions", "")
        items.append(f"{i}. [{cat}] {p['title']}\n   核心结论：{concl}")
    content = "\n".join(items)
    api = cfg["llm"]
    resp = requests.post(
        api["base_url"].rstrip("/") + "/chat/completions",
        headers={"Authorization": f"Bearer {api['api_key']}", "Content-Type": "application/json"},
        json={
            "model": api["model"],
            "messages": [
                {"role": "system",
                 "content": "你是土木工程组合结构领域的资深研究员。以下是今日筛选出的组合结构论文（标题+类别+核心结论）。"
                            "请写「今日论文总结与思考」，中英双语：\n"
                            "先中文：【总结】主题/材料/方法/趋势（300字）+【思考】热点/空白/交叉/启示（300字）；\n"
                            "再英文：【Summary】+【Reflection】各约250词。\n"
                            "用【】标注四个小标题，直接输出正文。"},
                {"role": "user", "content": content},
            ],
            "temperature": 0.5,
        },
        timeout=120,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def llm_comparison(cfg, papers):
    """横向对比各论文的研究逻辑与研究方法，形成对比总结。"""
    items = []
    for i, p in enumerate(papers, 1):
        cat = CATEGORIES.get(p.get("category"), "其他")
        methods = p.get("analysis", {}).get("methods", "")
        items.append(f"{i}. [{cat}] {p['title']}\n   方法：{methods}")
    content = "\n".join(items)
    api = cfg["llm"]
    resp = requests.post(
        api["base_url"].rstrip("/") + "/chat/completions",
        headers={"Authorization": f"Bearer {api['api_key']}", "Content-Type": "application/json"},
        json={
            "model": api["model"],
            "messages": [
                {"role": "system",
                 "content": "你是土木工程组合结构领域的研究方法专家。以下是今日论文（标题+类别+方法）。"
                            "请写「论文逻辑与方法对比总结」，中英双语：\n"
                            "先中文（600字左右）：横向对比研究逻辑（如何提出问题/验证）与研究方法（实验/数值/理论/机器学习）的异同，归纳共性主线与方法路线特点；\n"
                            "再英文（约400词）对应内容。\n"
                            "用【中文对比总结】【English Comparison】标注两段，直接输出正文。"},
                {"role": "user", "content": content},
            ],
            "temperature": 0.5,
        },
        timeout=120,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def main():
    cfg = fr.load_config()
    state = load_state()
    reported = set(state.get("reported_dois", []))
    journal_if = {norm(j["name"]): j for j in cfg["journals"]}

    if not os.path.exists(SD_JSON):
        print("未找到", SD_JSON)
        return 1
    with open(SD_JSON, "r", encoding="utf-8") as f:
        papers = json.load(f)

    valid = []
    for p in papers:
        if not p.get("doi") or not p.get("abstract"):
            continue
        jinfo = journal_if.get(norm(p.get("journal", "")))
        if not jinfo:
            continue
        p["if"] = jinfo.get("if", 0.0)
        valid.append(p)
    print(f"有效论文（DOI+摘要+Q1 期刊）: {len(valid)}")

    kept = []
    for i, p in enumerate(valid, 1):
        if p["doi"] in reported:
            continue
        p["cited"] = crossref_citations(p["doi"])
        try:
            cat, depth_ok = llm_classify(cfg, p["title"], p["abstract"])
        except Exception as e:
            print(f"  分类失败默认通过: {e}")
            cat, depth_ok = "1", True
        if cat and depth_ok:
            p["category"] = cat
            p["year"] = extract_year(p.get("date", ""))
            kept.append(p)
            print(f"  [{i}/{len(valid)}] [通过] 类别{cat}({CATEGORIES[cat]}) | 被引{p['cited']} | {p['title'][:36]}")
        else:
            print(f"  [{i}/{len(valid)}] [排除] 类{cat}/深度{depth_ok} {p['title'][:36]}")
        time.sleep(1)

    if not kept:
        fr.notify(cfg, "⚠️ 论文报告为空", "无通过质量校验的论文。")
        return 1

    for p in kept:
        p["score"] = p["if"] * math.log10(p["cited"] + 1)
    kept.sort(key=lambda x: (x["year"], x["score"]), reverse=True)

    selected = []
    for p in kept:
        if p["doi"] in reported:
            continue
        selected.append(p)
        if len(selected) >= int(cfg["max_papers"]):
            break

    if not selected:
        msg = "本日无新增论文（已全部推送过），明日再试。"
        print(msg)
        fr.notify(cfg, "ℹ️ 组合结构论文报告：无新增", msg)
        return 0

    for p in selected:
        try:
            p["analysis"] = llm_analysis_with_intro(cfg, p["abstract"], p.get("introduction", ""))
        except Exception:
            p["analysis"] = fr.heuristic_analysis(p["abstract"])

    try:
        summary = llm_summary(cfg, selected)
    except Exception as e:
        print(f"  总结生成失败: {e}")
        summary = ""

    try:
        comparison = llm_comparison(cfg, selected)
    except Exception as e:
        print(f"  对比总结生成失败: {e}")
        comparison = ""

    today = datetime.date.today().strftime("%Y-%m-%d")
    out_dir = cfg["output_dir"]
    os.makedirs(out_dir, exist_ok=True)
    warn = ("数据源：ScienceDirect（浏览器自动化）+ Crossref 被引量；近 5 年、Q1 顶刊白名单；"
            "按 年份(优先今年) → 影响因子×log10(被引+1) 排序；已做 8 大类主题匹配 + 深度校验 + 去重。")
    md_path = os.path.join(out_dir, f"{today}+论文报告.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(build_markdown_sd(selected, today, warn, summary, comparison))

    # 调用 Claude Code 批量摄入（继承 E:\project_wiki\CLAUDE.md 的 ingest 规范）
    import subprocess
    ingest_prompt = ("批量 ingest E:/project_wiki/raw/reports_literature，无需逐步确认："
                     "把报告中每篇论文的信息、摘要、引言逐篇摄入知识库"
                     "（每篇建 source 页、更新 concept/entity 页）；"
                     "完成后执行 lint 健康检查；最后更新 index.md 和 log.md。完成后简要汇报。")
    try:
        r = subprocess.run(
            f'claude -p "{ingest_prompt}" --dangerously-skip-permissions',
            cwd="E:/project_wiki", shell=True, capture_output=True, text=True, timeout=3600,
        )
        print(f"  claude ingest: exit={r.returncode}")
        if r.stdout:
            print("  [claude 输出末尾]")
            print(r.stdout[-800:])
    except Exception as e:
        print(f"  claude ingest 失败: {e}")

    for p in selected:
        reported.add(p["doi"])
    state["reported_dois"] = list(reported)
    state["last_run"] = today
    save_state(state)

    cat_cnt = {}
    for p in selected:
        k = CATEGORIES.get(p["category"], p["category"])
        cat_cnt[k] = cat_cnt.get(k, 0) + 1
    cat_str = "，".join(f"{k} {v}" for k, v in cat_cnt.items())

    fr.notify(cfg, f"📄 {today} 组合结构论文报告已生成",
              f"共 {len(selected)} 篇（去重后新增）\n类别分布：{cat_str}\n\n文件：\n{md_path}")
    print(f"完成：{len(selected)} 篇 -> {md_path}")
    return 0


def build_markdown_sd(papers, today, warn, summary="", comparison=""):
    lines = [
        "---",
        "type: source",
        f"title: \"{today} 组合结构论文报告\"",
        f"date: {today}",
        "tags: [组合结构, 论文报告, 每日论文猎人]",
        "processed: false",
        "---",
        "",
        f"# {today} 组合结构论文报告", "",
    ]
    lines.append(f"**共 {len(papers)} 篇**（ScienceDirect 抓取，优先今年 → 影响因子×log10(被引+1) 排序）")
    lines.append(f"**生成时间**: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    lines.append(f"> {warn}")
    lines.append("")
    for i, p in enumerate(papers, 1):
        cat = CATEGORIES.get(p.get("category"), "其他")
        lines.append(f"## {i}. {p['title']}")
        lines.append("")
        lines.append(f"- **类别**: {cat}")
        lines.append(f"- **期刊**: {p['journal']}（影响因子 {p['if']}）")
        lines.append(f"- **作者**: {', '.join(p.get('authors', []) or []) or '—'}")
        lines.append(f"- **发表日期**: {p.get('date', '—')}")
        lines.append(f"- **被引次数**: {p['cited']}")
        lines.append(f"- **DOI**: [{p['doi']}](https://doi.org/{p['doi']})")
        lines.append("")
        lines.append("### 摘要")
        lines.append("")
        lines.append(p["abstract"])
        if p.get("introduction"):
            lines.append("")
            lines.append("### 引言")
            lines.append("")
            lines.append(p["introduction"])
        lines.append("")
        lines.append("### 研究逻辑与方法")
        lines.append("")
        lines.append(p["analysis"]["objective"] + "\n\n" + p["analysis"]["methods"])
        if p["analysis"].get("objective_en"):
            lines.append("")
            lines.append("**Research Logic & Methods (English)**")
            lines.append("")
            lines.append(p["analysis"]["objective_en"] + "\n\n" + p["analysis"]["methods_en"])
        lines.append("")
        lines.append("### 核心结论")
        lines.append("")
        lines.append(p["analysis"]["conclusions"])
        if p["analysis"].get("conclusions_en"):
            lines.append("")
            lines.append("**Core Conclusions (English)**")
            lines.append("")
            lines.append(p["analysis"]["conclusions_en"])
        lines.append("")
        lines.append("---")
        lines.append("")
    if summary:
        lines.append("## 今日论文总结与思考")
        lines.append("")
        lines.append(summary)
        lines.append("")
    if comparison:
        lines.append("## 论文逻辑与方法对比总结")
        lines.append("")
        lines.append(comparison)
        lines.append("")
    return "\n".join(lines)


def build_docx_sd(papers, today, warn, summary, comparison, path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    doc = Document()
    for style_name, size in [("Normal", 11), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), "微软雅黑")
    doc.add_heading(f"{today} 组合结构论文报告", level=1)
    doc.add_paragraph(f"共 {len(papers)} 篇（优先今年 → 影响因子×log10(被引+1) 排序）")
    doc.add_paragraph(f"生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"⚠️ {warn}")
    for i, p in enumerate(papers, 1):
        cat = CATEGORIES.get(p.get("category"), "其他")
        doc.add_heading(f"{i}. {p['title']}", level=2)
        for label, val in [
            ("类别", cat),
            ("期刊", f"{p['journal']}（影响因子 {p['if']}）"),
            ("作者", ", ".join(p.get("authors", []) or []) or "—"),
            ("发表日期", p.get("date", "—")),
            ("被引次数", str(p["cited"])),
            ("DOI", f"https://doi.org/{p['doi']}"),
        ]:
            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            para.add_run(val)
        doc.add_heading("摘要", level=3)
        doc.add_paragraph(p["abstract"])
        if p.get("introduction"):
            doc.add_heading("引言", level=3)
            doc.add_paragraph(p["introduction"])
        doc.add_heading("研究逻辑与方法", level=3)
        doc.add_paragraph(p["analysis"]["objective"])
        doc.add_paragraph(p["analysis"]["methods"])
        if p["analysis"].get("objective_en"):
            doc.add_heading("Research Logic & Methods (English)", level=3)
            doc.add_paragraph(p["analysis"]["objective_en"])
            doc.add_paragraph(p["analysis"]["methods_en"])
        doc.add_heading("核心结论", level=3)
        doc.add_paragraph(p["analysis"]["conclusions"])
        if p["analysis"].get("conclusions_en"):
            doc.add_heading("Core Conclusions (English)", level=3)
            doc.add_paragraph(p["analysis"]["conclusions_en"])
    if summary:
        doc.add_heading("今日论文总结与思考", level=2)
        doc.add_paragraph(summary)
    if comparison:
        doc.add_heading("论文逻辑与方法对比总结", level=2)
        doc.add_paragraph(comparison)
    doc.save(path)


if __name__ == "__main__":
    sys.exit(main())

